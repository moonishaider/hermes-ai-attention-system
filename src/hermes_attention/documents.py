"""Private, bounded document operations. Input text never conveys tool authority.

Only fixed extraction and generation operations run; this is not a Python sandbox.
The caller supplies an already-authorized conversation and selected file/bytes.
"""
from __future__ import annotations

import csv
import hashlib
import io
import json
import os
from pathlib import Path, PurePosixPath
import re
import stat
import threading
import contextlib
import fcntl
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from datetime import datetime, timezone, timedelta
from xml.sax.saxutils import escape

MAX_BYTES = 25 * 1024 * 1024
MAX_EXPANDED = 100 * 1024 * 1024
MAX_UNITS = 100_000
MAX_TEXT = 2_000_000
TYPES = {'.txt':'text/plain', '.md':'text/markdown', '.csv':'text/csv', '.pdf':'application/pdf', '.docx':'application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.xlsx':'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', '.png':'image/png', '.jpg':'image/jpeg', '.jpeg':'image/jpeg', '.webp':'image/webp'}
_LOCK = threading.RLock()
_DEPTH = threading.local()


@contextlib.contextmanager
def _locked(root):
    # Locks survive separate CLI invocations; nested same-thread reads are safe.
    with _LOCK:
        depths = getattr(_DEPTH, 'roots', {})
        key = str(root)
        if depths.get(key, 0):
            yield
            return
        lock_path = _no_symlinks(Path(root) / '.documents.lock')
        fd = os.open(lock_path, os.O_RDWR | os.O_CREAT | getattr(os, 'O_NOFOLLOW', 0), 0o600)
        try:
            fcntl.flock(fd, fcntl.LOCK_EX)
            depths[key] = 1
            _DEPTH.roots = depths
            yield
        finally:
            depths.pop(key, None)
            fcntl.flock(fd, fcntl.LOCK_UN)
            os.close(fd)


def _now():
    return datetime.now(timezone.utc).isoformat()


def _identifier(value):
    if not isinstance(value, str) or not re.fullmatch(r'[A-Za-z0-9_-]{1,160}', value):
        raise ValueError('Invalid document or conversation identity')
    return value


def _safe_name(value):
    if not isinstance(value, str) or not value or len(value) > 240 or '\\' in value or '/' in value or any(ord(c) < 32 for c in value) or value in {'.', '..'}:
        raise ValueError('A plain filename is required')
    return value


def _no_symlinks(path):
    path = Path(path).absolute()
    if any(p.is_symlink() for p in [path, *path.parents]):
        raise ValueError('Symlink paths are not permitted')
    return path


def _private_write(path, data):
    _no_symlinks(path)
    fd = os.open(path, os.O_WRONLY | os.O_CREAT | os.O_EXCL | getattr(os, 'O_NOFOLLOW', 0), 0o600)
    with os.fdopen(fd, 'wb') as stream:
        stream.write(data)
        stream.flush()
        os.fsync(stream.fileno())


def _check_zip(data, extension):
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        members = archive.infolist()
        if len(members) > 5000 or sum(x.file_size for x in members) > MAX_EXPANDED:
            raise ValueError('Document expansion limit exceeded')
        names = set()
        for item in members:
            name = item.filename
            if name in names or '\\' in name or PurePosixPath(name).is_absolute() or '..' in PurePosixPath(name).parts or stat.S_ISLNK(item.external_attr >> 16):
                raise ValueError('Unsafe archive member')
            names.add(name)
            if item.flag_bits & 1 or item.file_size > MAX_BYTES or item.file_size > max(1, item.compress_size) * 300:
                raise ValueError('Encrypted or excessive archive member')
            if any(t in name.lower() for t in ('vbaproject', '/embeddings/', 'activex')):
                raise ValueError('Macros and embedded executable objects are unsupported')
        required = 'word/document.xml' if extension == '.docx' else 'xl/workbook.xml'
        if required not in names:
            raise ValueError('Document content does not match its extension')


class DocumentWorkspace:
    """Manifest lives under an owned private directory; no provider/network calls."""
    def __init__(self, root):
        self.root = _no_symlinks(root)
        self.root.mkdir(mode=0o700, parents=True, exist_ok=True)
        self.manifest = self.root / 'documents.json'

    def _read(self):
        _no_symlinks(self.manifest)
        return json.loads(self.manifest.read_text()) if self.manifest.exists() else {'version': 1, 'documents': {}}

    def _save(self, manifest):
        path = self.root / ('.manifest-' + uuid.uuid4().hex)
        _private_write(path, json.dumps(manifest, ensure_ascii=False).encode())
        os.replace(path, self.manifest)

    def list(self, conversation_id, include_forgotten=False):
        _identifier(conversation_id)
        with _locked(self.root):
            return [r for r in self._read()['documents'].values() if r['conversation_id'] == conversation_id and (include_forgotten or (r['retention_state'] == 'active' and (not r.get('expires_at') or r['expires_at'] > _now())))]

    def get(self, document_id, conversation_id):
        with _locked(self.root):
            record = self._read()['documents'].get(_identifier(document_id))
            if not record or record['conversation_id'] != _identifier(conversation_id) or record['retention_state'] != 'active' or (record.get('expires_at') and record['expires_at'] <= _now()):
                raise ValueError('Document unavailable in this conversation')
            return record

    def path(self, document_id, conversation_id):
        record = self.get(document_id, conversation_id)
        path = _no_symlinks(self.root / record['storage_name'])
        if path.parent != self.root or not path.is_file():
            raise ValueError('Document file unavailable')
        return path

    def ingest_file(self, path, *, conversation_id, turn_id='', retention='conversation', source='selected-file', parent_id=None):
        path = _no_symlinks(path)
        if not path.is_file() or path.stat().st_size > MAX_BYTES:
            raise ValueError('Select one regular file up to 25 MiB')
        fd = os.open(path, os.O_RDONLY | getattr(os, 'O_NOFOLLOW', 0))
        with os.fdopen(fd, 'rb') as stream:
            data = stream.read(MAX_BYTES + 1)
        return self.ingest_bytes(data, name=path.name, conversation_id=conversation_id, turn_id=turn_id, retention=retention, source=source, parent_id=parent_id)

    def ingest_bytes(self, data, *, name, conversation_id, turn_id='', retention='conversation', source='attachment', parent_id=None):
        _identifier(conversation_id)
        if turn_id:
            _identifier(turn_id)
        _safe_name(name)
        if retention not in {'conversation', 'ephemeral'}:
            raise ValueError('Unknown retention choice')
        extension = Path(name).suffix.lower()
        if extension not in TYPES or not data or len(data) > MAX_BYTES:
            raise ValueError('Unsupported, empty, or oversized attachment')
        if extension in {'.xlsx', '.docx'}:
            _check_zip(data, extension)
        if extension == '.pdf' and not data.startswith(b'%PDF-'):
            raise ValueError('Content is not a PDF')
        digest = hashlib.sha256(data).hexdigest()
        with _locked(self.root):
            manifest = self._read()
            parent = self.get(parent_id, conversation_id) if parent_id else None
            document_id = 'doc_' + uuid.uuid4().hex
            storage_name = document_id + extension
            _private_write(self.root / storage_name, data)
            record = {'id': document_id, 'original_name':name, 'display_name':name, 'mime':TYPES[extension], 'sha256':digest, 'bytes':len(data), 'source':source, 'conversation_id':conversation_id, 'turn_id':turn_id, 'version':parent['version'] + 1 if parent else 1, 'parent_id':parent_id, 'retention':retention, 'retention_state':'active', 'storage_name':storage_name, 'created_at':_now(), 'expires_at':(datetime.now(timezone.utc)+timedelta(minutes=15)).isoformat() if retention=='ephemeral' else None, 'extraction_status':'pending', 'warnings':[], 'authority':'untrusted-source-data'}
            manifest['documents'][document_id] = record
            self._save(manifest)
        return self.extract(document_id, conversation_id)

    def extract(self, document_id, conversation_id):
        record = self.get(document_id, conversation_id)
        path = self.path(document_id, conversation_id)
        units, warnings = [], []
        chars = 0
        def add(locator, text, **extra):
            nonlocal chars
            text = str(text)
            chars += len(text)
            if chars > MAX_TEXT or len(units) >= MAX_UNITS:
                raise ValueError('Extraction limit reached; document is incomplete')
            units.append({'citation':f"{document_id}:{locator}", 'locator':locator, 'text':text, **extra})
        try:
            ext = path.suffix
            if ext in {'.txt', '.md'}:
                for i, line in enumerate(path.read_text(encoding='utf-8-sig').splitlines(), 1):
                    add(f'line:{i}', line)
            elif ext == '.csv':
                with path.open(encoding='utf-8-sig', newline='') as stream:
                    for i, row in enumerate(csv.reader(stream), 1):
                        add(f'row:{i}', ' | '.join(row), cells=row)
            elif ext == '.docx':
                from docx import Document
                doc = Document(path)
                for i, paragraph in enumerate(doc.paragraphs, 1):
                    add(f'paragraph:{i}', paragraph.text)
                for t, table in enumerate(doc.tables, 1):
                    for i, row in enumerate(table.rows, 1):
                        add(f'table:{t}:row:{i}', ' | '.join(c.text for c in row.cells), cells=[c.text for c in row.cells])
                if doc.inline_shapes:
                    warnings.append('Embedded images require selective vision; paragraph and table text extracted.')
            elif ext == '.xlsx':
                from openpyxl import load_workbook
                formula_book = load_workbook(path, read_only=True, data_only=False, keep_links=False)
                value_book = load_workbook(path, read_only=True, data_only=True, keep_links=False)
                try:
                    for sheet in formula_book:
                        if sheet.max_row * sheet.max_column > MAX_UNITS:
                            raise ValueError('Workbook cell limit exceeded')
                        values = value_book[sheet.title]
                        for row in sheet:
                            for cell in row:
                                if cell.value is not None:
                                    formula = cell.value if cell.data_type == 'f' else None
                                    cached = values[cell.coordinate].value if formula else cell.value
                                    add(f'sheet:{sheet.title}:cell:{cell.coordinate}', cached if cached is not None else '', formula=formula, cached_value=cached, sheet_state=sheet.sheet_state)
                                    if formula and cached is None:
                                        warnings.append(f'{sheet.title}!{cell.coordinate}: formula has no cached value; not calculated')
                finally:
                    formula_book.close(); value_book.close()
            elif ext == '.pdf':
                from pypdf import PdfReader
                reader = PdfReader(path)
                if reader.is_encrypted:
                    raise ValueError('Encrypted PDF requires an owner-decrypted copy')
                if len(reader.pages) > 500:
                    raise ValueError('PDF page limit exceeded')
                for i, page in enumerate(reader.pages, 1):
                    text = page.extract_text() or ''
                    add(f'page:{i}', text, rotation=page.rotation)
                    if not text.strip():
                        warnings.append(f'Page {i} needs OCR or approved vision; native text unavailable')
            else:
                from PIL import Image
                with Image.open(path) as img:
                    if img.width * img.height > 25_000_000:
                        raise ValueError('Image pixel limit exceeded')
                    expected = {'.png':'PNG', '.jpg':'JPEG', '.jpeg':'JPEG', '.webp':'WEBP'}[ext]
                    if img.format != expected:
                        raise ValueError('Image content does not match its extension')
                    img.verify()
                    add('image:1', '', width=img.width, height=img.height)
                warnings.append('Image retained for approved vision; no text extraction has been claimed')
            status = 'needs_vision' if any('needs OCR' in w or 'require selective vision' in w or 'retained for approved vision' in w for w in warnings) else 'complete_with_warnings' if warnings else 'complete'
        except (Exception,) as exc:
            status = 'failed'
            warnings.append(f'{type(exc).__name__}: {str(exc)[:240]}')
        with _locked(self.root):
            manifest = self._read()
            current = manifest['documents'][document_id]
            current.update(extraction_status=status, units=units, warnings=list(dict.fromkeys(warnings)), extracted_at=_now(), extraction_complete=status in {'complete', 'complete_with_warnings'})
            self._save(manifest)
            return current

    def ocr(self, document_id, conversation_id, *, max_pages=12):
        """Reviewed local OCR operation. No model/provider calls or source links.

        Explicitly bounded to 12 selected textless pages per call, 30s per rotation.
        Tesseract executes only against copied images with no inherited credentials.
        Returned OCR remains source data and includes its method/rotation/confidence.
        """
        if not 1 <= max_pages <= 12:
            raise ValueError('OCR page bound must be between 1 and 12')
        record = self.get(document_id, conversation_id)
        path = self.path(document_id, conversation_id)
        if path.suffix not in {'.pdf','.png','.jpg','.jpeg','.webp'}:
            raise ValueError('Local OCR supports PDF pages and selected images; extract embedded DOCX images separately')
        binary = shutil.which('tesseract')
        if not binary:
            return {**record, 'ocr_status':'unavailable', 'ocr_warning':'Local Tesseract is unavailable; approved vision adapter required'}
        from PIL import Image, ImageOps
        units = list(record.get('units', []))
        candidates = [u for u in units if not u['text'].strip() and (u['locator'].startswith('page:') or u['locator']=='image:1')]
        completed = []
        with tempfile.TemporaryDirectory(prefix='.ocr-', dir=self.root) as temporary:
            working = Path(temporary)
            pdf = None
            try:
                if path.suffix == '.pdf':
                    import pypdfium2
                    pdf = pypdfium2.PdfDocument(str(path))
                for unit in candidates[:max_pages]:
                    if pdf is not None:
                        page = pdf[int(unit['locator'].split(':')[1])-1]
                        if page.get_width()*page.get_height()*4 > 25_000_000:
                            raise ValueError('Rendered page pixel bound exceeded')
                        bitmap = page.render(scale=2)
                        image = bitmap.to_pil().copy()
                        bitmap.close(); page.close()
                    else:
                        with Image.open(path) as source:
                            image = ImageOps.exif_transpose(source).convert('RGB')
                    best = ('', -1.0, 0)
                    for rotation in (0, 90, 180, 270):
                        selected = image.rotate(rotation, expand=True)
                        image_path = working / 'selected.png'
                        selected.save(image_path)
                        result = subprocess.run([binary, str(image_path), 'stdout', '--psm', '6', 'tsv'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=30, env={'PATH':'/usr/bin:/bin','LANG':'en_US.UTF-8','OMP_THREAD_LIMIT':'1'}, cwd=working, check=True)
                        if len(result.stdout)>MAX_TEXT:
                            raise ValueError('OCR output bound exceeded')
                        rows=list(csv.DictReader(io.StringIO(result.stdout.decode('utf-8')),delimiter='\t'))
                        words=[r for r in rows if r.get('text','').strip() and float(r.get('conf',-1))>=0]
                        text=' '.join(r['text'] for r in words)
                        confidence=sum(float(r['conf']) for r in words)/len(words) if words else 0
                        if confidence>best[1]: best=(text,confidence,rotation)
                        if confidence>=92: break
                    unit.update(text=best[0], extraction_method='local-tesseract', ocr_confidence=round(best[1],2), ocr_rotation=best[2])
                    completed.append(unit['locator'])
            finally:
                if pdf is not None: pdf.close()
        remaining=[u['locator'] for u in units if not u['text'].strip()]
        with _locked(self.root):
            manifest=self._read(); current=manifest['documents'][document_id]
            if current['retention_state']!='active': raise ValueError('Attachment revoked during extraction')
            current.update(units=units,ocr_status='complete' if not remaining else 'partial',ocr_completed=completed,ocr_remaining=remaining,extraction_status='complete_with_warnings' if not remaining else 'needs_vision',extraction_complete=not remaining)
            current['warnings']=['OCR is machine extracted and may need visual verification.'] + ([f'{len(remaining)} pages/images still need extraction or are blank.'] if remaining else [])
            self._save(manifest)
            return current

    def forget(self, document_id, conversation_id):
        self.get(document_id, conversation_id)
        with _locked(self.root):
            manifest = self._read()
            record = manifest['documents'][document_id]
            record.update(retention_state='forgotten', forgotten_at=_now())
            self._save(manifest)
            return {'id':document_id, 'retrieval_revoked':True, 'bytes_deleted':False, 'recoverable':True, 'explanation':'Retrieval and opening revoked. Private copy remains recoverable; original file unchanged.'}

    def restore(self, document_id, conversation_id):
        with _locked(self.root):
            manifest = self._read()
            record = manifest['documents'].get(_identifier(document_id))
            if not record or record['conversation_id'] != _identifier(conversation_id):
                raise ValueError('Document unavailable')
            record['retention_state'] = 'active'
            self._save(manifest)
            return record

    def generate(self, *, conversation_id, format, title, sections=(), tables=(), source_ids=(), parent_id=None, turn_id=''):
        """Tables: {name, headers, rows}. Cell values are data; formulas are not accepted."""
        if format not in {'txt','md','csv','xlsx','docx','pdf'}:
            raise ValueError('Unsupported output format')
        if not isinstance(title, str) or not title.strip():
            raise ValueError('A document title is required')
        if len(str(title)) > 200 or len(json.dumps([sections, tables], default=str)) > MAX_TEXT:
            raise ValueError('Output limit exceeded')
        sources = [self.get(i, conversation_id) for i in source_ids]
        sections = list(sections) + ([{'heading':'Sources', 'text':'\n'.join(f"{s['display_name']} | {s['id']} | SHA256 {s['sha256']}" for s in sources)}] if sources else [])
        for table in tables:
            if not table['headers'] or len(table['headers']) > 30 or any(len(row) != len(table['headers']) for row in table['rows']):
                raise ValueError('Invalid table dimensions')
        out = io.BytesIO()
        if format in {'txt','md'}:
            out.write((title + '\n\n' + '\n\n'.join(str(s.get('heading',''))+'\n'+str(s.get('text','')) for s in sections) + '\n\n' + '\n\n'.join(t['name']+'\n'+'\n'.join(' | '.join(map(str,r)) for r in [t['headers'], *t['rows']]) for t in tables)).encode())
        elif format == 'csv':
            if len(tables) != 1:
                raise ValueError('CSV requires exactly one table')
            text = io.StringIO(newline=''); writer = csv.writer(text)
            for row in [tables[0]['headers'], *tables[0]['rows']]:
                writer.writerow([_spreadsheet_text(v) for v in row])
            out.write(text.getvalue().encode('utf-8-sig'))
        elif format == 'xlsx':
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            book = Workbook(); book.remove(book.active)
            for index, table in enumerate(tables or [{'name':'Report','headers':['Section','Text'],'rows':[[s.get('heading',''), s.get('text','')] for s in sections]}]):
                sheet = book.create_sheet(re.sub(r'[\\/*?:\[\]]',' ',table['name'])[:26] + f' {index+1}')
                for row in [table['headers'], *table['rows']]:
                    sheet.append([_spreadsheet_text(v) for v in row])
                from decimal import Decimal
                for row in sheet.iter_rows(min_row=2):
                    for cell in row:
                        if isinstance(cell.value,Decimal):cell.number_format='#,##0.00########'
                sheet.freeze_panes = 'A2'; sheet.auto_filter.ref = sheet.dimensions
                for cell in sheet[1]:
                    cell.font = Font(bold=True, color='FFFFFF'); cell.fill = PatternFill('solid', fgColor='183348')
                for col in sheet.columns:
                    sheet.column_dimensions[col[0].column_letter].width = min(55, max(12, max(len(str(c.value or '')) for c in col)+2))
            if sections and tables:
                sheet = book.create_sheet('Notes and sources')
                for s in sections:
                    for index,line in enumerate(str(s.get('text','')).splitlines() or ['']):
                        sheet.append([s.get('heading','') if index==0 else '',_spreadsheet_text(line)])
                sheet.column_dimensions['A'].width=32;sheet.column_dimensions['B'].width=100
                for row in sheet:
                    for cell in row:cell.alignment=Alignment(wrap_text=True,vertical='top')
                    sheet.row_dimensions[row[0].row].height=max(30,min(150,15*(1+len(str(row[1].value or ''))//100)))
            book.save(out)
        elif format == 'docx':
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            doc = Document(); doc.add_heading(title, 0)
            for style in doc.styles:
                for border in list(style.element.iter(qn('w:pBdr'))):
                    border.getparent().remove(border)
            for key in ('Title','Heading 1','Heading 2'):
                doc.styles[key].font.color.rgb = RGBColor(0,0,0)
            doc.styles['Normal'].font.size = Pt(10)
            doc.styles['Normal'].font.name = 'Arial'
            for s in sections:
                if s.get('heading'): doc.add_heading(str(s['heading']), 1)
                for paragraph in str(s.get('text','')).split('\n'): doc.add_paragraph(paragraph)
            for table in tables:
                doc.add_heading(table['name'], 1)
                grid = doc.add_table(rows=1, cols=len(table['headers'])); grid.style = 'Light Shading Accent 1'
                repeat_header=OxmlElement('w:tblHeader')
                grid.rows[0]._tr.get_or_add_trPr().append(repeat_header)
                for cell, value in zip(grid.rows[0].cells,table['headers']): cell.text = str(value)
                for row in table['rows']:
                    for cell, value in zip(grid.add_row().cells,row): cell.text = str(value)
            for section in doc.sections: section.left_margin = section.right_margin = Inches(.75)
            doc.save(out)
        else:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            styles = getSampleStyleSheet(); story = [Paragraph(escape(title), styles['Title'])]
            for s in sections:
                if s.get('heading'): story.append(Paragraph(escape(str(s['heading'])), styles['Heading2']))
                for line in str(s.get('text','')).split('\n'):
                    story += [Paragraph(escape(line), styles['BodyText']), Spacer(1,5)]
            for table in tables:
                story.append(Paragraph(escape(table['name']), styles['Heading2']))
                if len(table['headers']) > 8: raise ValueError('PDF tables support up to eight columns; use XLSX for wide data')
                rows = [[Paragraph(escape(str(v)),styles['BodyText']) for v in r] for r in [table['headers'],*table['rows']]]
                grid = Table(rows, repeatRows=1, colWidths=[468/len(table['headers'])]*len(table['headers']))
                grid.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#e9f1f5')),('VALIGN',(0,0),(-1,-1),'TOP'),('BOTTOMPADDING',(0,0),(-1,-1),6)])); story.append(grid)
            SimpleDocTemplate(out, leftMargin=72,rightMargin=72).build(story)
        name = re.sub(r'[^\w .-]', '_', title)[:120].strip(' .') or 'Report'
        return self.ingest_bytes(out.getvalue(),name=name+'.'+format,conversation_id=conversation_id,turn_id=turn_id,source='generated-fixed-operation',parent_id=parent_id)


def _spreadsheet_text(value):
    from decimal import Decimal
    if isinstance(value,Decimal) and len(value.normalize().as_tuple().digits)>15:
        return format(value,'f')  # Preserve exact high-precision values as editable text.
    if isinstance(value, str) and value.lstrip().startswith(('=', '+', '-', '@', '\t', '\r')):
        return "'" + value
    return value
